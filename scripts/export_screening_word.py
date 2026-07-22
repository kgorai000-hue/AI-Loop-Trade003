from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]
XLSX = PROJECT_ROOT / "data/reports/screening_H1_2000bars_20260711_234506.xlsx"
DOCX = XLSX.with_suffix(".docx")


def main() -> None:
    summary = pd.read_excel(XLSX, sheet_name="Summary")
    detail = pd.read_excel(XLSX, sheet_name="StrategyDetail")

    doc = Document()
    doc.add_heading("MT5_loop 全銘柄スクリーニング報告書", 0)
    doc.add_paragraph("条件: H1 / 2000本 / Monte Carlo 100回 / 12銘柄順次実行")
    doc.add_paragraph("生成日時(UTC): 2026-07-12 00:04")

    doc.add_heading("1. 総合結果", level=1)
    doc.add_paragraph("品質ゲート合格(1戦略以上): 0 / 12 銘柄")
    doc.add_paragraph("エラー: 0件")
    doc.add_paragraph(
        "注記: スクリーニングモードのため厳しいゲート基準を満たさない結果。"
        "有望銘柄は9000本以上で再検証を推奨。"
    )

    doc.add_heading("2. 銘柄別サマリー", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, t in enumerate(["銘柄", "ゲート合格", "最良戦略", "Sharpe", "WF Sharpe", "備考"]):
        hdr[i].text = t
    for _, r in summary.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["symbol"])
        row[1].text = f"{int(r['n_pass'])}/3"
        row[2].text = str(r["best_strategy"])
        row[3].text = str(r["best_sharpe"])
        row[4].text = str(r["best_wf_test_sharpe"])
        if r["best_sharpe"] == 0 and r["best_wf_test_sharpe"] == 0:
            row[5].text = "シグナルなし"
        else:
            row[5].text = "要再検証"

    doc.add_heading("3. 相対的に有望な銘柄（参考）", level=1)
    best = (
        detail.groupby("symbol")
        .agg({"wf_avg_test_sharpe": "max", "total_return_pct": "max", "trades": "max"})
        .reset_index()
        .sort_values("wf_avg_test_sharpe", ascending=False)
        .head(5)
    )
    for _, r in best.iterrows():
        doc.add_paragraph(
            f"- {r['symbol']}: WF test Sharpe {r['wf_avg_test_sharpe']:.3f}, "
            f"max return {r['total_return_pct']:.1f}%, trades {int(r['trades'])}"
        )

    doc.add_heading("4. 主な不合格理由", level=1)
    doc.add_paragraph("- OOS performance (test/train ratio < 0.5)")
    doc.add_paragraph("- Monte Carlo P5 マイナス")
    doc.add_paragraph("- feature_score 取引0件: Japan225, GOLD, SILVER, WTI")
    doc.add_paragraph("- Parameter stability (trend戦略)")

    doc.add_heading("5. 詳細データ", level=1)
    doc.add_paragraph(f"Excel: {XLSX.name}")

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
