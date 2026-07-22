from __future__ import annotations

import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

import pandas as pd
from docx import Document

from scripts.screening_params import BEST_MR_RSI_BY_SYMBOL, rsi_band_label

BASE = PROJECT_ROOT / "data/reports"

BEST_RSI = {symbol: rsi_band_label(symbol) for symbol in BEST_MR_RSI_BY_SYMBOL}

RSI_FILES = {
    "30/70": BASE / "screening_M30_2000bars_nogate_20260715_133134.xlsx",
    "25/75": BASE / "screening_M30_2000bars_nogate_rsi2575_20260715_201256.xlsx",
    "35/65": BASE / "screening_M30_2000bars_nogate_rsi3565_20260715_204330.xlsx",
}

BB20 = BASE / "screening_M30_2000bars_nogate_rsibest_bb2080_20260715_213606.xlsx"
BB10 = BASE / "screening_M30_2000bars_nogate_rsibest_bb1090_20260715_215954.xlsx"
DOCX = BB10.with_name(BB10.stem + "_report.docx")


def mr_rows(path: Path) -> pd.DataFrame:
    d = pd.read_excel(path, sheet_name="StrategyDetail")
    return d[d["strategy"] == "mean_reversion"].copy()


def main() -> None:
    bb10_path = BB10
    if not bb10_path.exists():
        matches = sorted(BASE.glob("screening_M30_2000bars_nogate_rsibest_bb1090_*.xlsx"))
        if not matches:
            raise FileNotFoundError("BB 0.10/0.90 report not found")
        bb10_path = matches[-1]

    baselines: list[dict] = []
    for symbol, band in BEST_RSI.items():
        row = mr_rows(RSI_FILES[band])
        r = row[row["symbol"] == symbol].iloc[0]
        baselines.append(
            {
                "symbol": symbol,
                "rsi": band,
                "ret_15": float(r["total_return_pct"]),
                "sharpe_15": float(r["sharpe"]),
                "mdd_15": float(r["max_drawdown_pct"]),
            }
        )
    base_df = pd.DataFrame(baselines)

    def side(path: Path, suffix: str) -> pd.DataFrame:
        m = mr_rows(path)[
            ["symbol", "total_return_pct", "sharpe", "max_drawdown_pct", "trades", "mc_prob_positive_pct"]
        ].rename(
            columns={
                "total_return_pct": f"ret_{suffix}",
                "sharpe": f"sharpe_{suffix}",
                "max_drawdown_pct": f"mdd_{suffix}",
                "trades": f"trades_{suffix}",
                "mc_prob_positive_pct": f"mc_{suffix}",
            }
        )
        return m

    cmp = (
        base_df.merge(side(BB20, "20"), on="symbol")
        .merge(side(bb10_path, "10"), on="symbol")
    )
    cmp["best_bb"] = cmp.apply(
        lambda r: max(
            [("0.15/0.85", r["ret_15"]), ("0.20/0.80", r["ret_20"]), ("0.10/0.90", r["ret_10"])],
            key=lambda x: x[1],
        )[0],
        axis=1,
    )
    cmp = cmp.sort_values("ret_10", ascending=False)

    print("=== Best RSI fixed | BB 0.15 / 0.20 / 0.10 ===")
    print(
        cmp[
            ["symbol", "rsi", "ret_15", "ret_20", "ret_10", "best_bb"]
        ].to_string(index=False)
    )
    print()
    print(cmp["best_bb"].value_counts().to_string())

    doc = Document()
    doc.add_heading("M30 BB entry 0.10/0.90 全銘柄テスト", 0)
    doc.add_paragraph("条件: M30 / 2000本 / MC100 / ゲートOFF")
    doc.add_paragraph("RSI: 銘柄別最良値に固定")
    doc.add_paragraph("変更: bb_entry 0.10/0.90（0.15・0.20と比較）")

    doc.add_heading("1. mean_reversion 3方式比較", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for i, t in enumerate(["銘柄", "RSI", "0.15", "0.20", "0.10", "最良BB"]):
        table.rows[0].cells[i].text = t
    for _, r in cmp.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["symbol"])
        row[1].text = str(r["rsi"])
        row[2].text = f"{r['ret_15']:.1f}%"
        row[3].text = f"{r['ret_20']:.1f}%"
        row[4].text = f"{r['ret_10']:.1f}%"
        row[5].text = str(r["best_bb"])

    doc.add_heading("2. 所見", level=1)
    for band, count in cmp["best_bb"].value_counts().items():
        doc.add_paragraph(f"- 最良 {band}: {int(count)} 銘柄")
    doc.add_paragraph(f"Excel: {bb10_path.name}")
    out = bb10_path.with_name(bb10_path.stem + "_report.docx")
    doc.save(out)
    print("Word:", out)


if __name__ == "__main__":
    main()
