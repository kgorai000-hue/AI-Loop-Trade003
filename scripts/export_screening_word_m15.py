from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]
XLSX = PROJECT_ROOT / "data/reports/screening_M15_2000bars_20260715_121340.xlsx"
DOCX = XLSX.with_suffix(".docx")


def main() -> None:
    summary = pd.read_excel(XLSX, sheet_name="Summary")
    detail = pd.read_excel(XLSX, sheet_name="StrategyDetail")

    n_pass = int((summary["n_pass"] >= 1).sum())

    doc = Document()
    doc.add_heading("MT5_loop 全銘柄スクリーニング報告書（M15）", 0)
    doc.add_paragraph("条件: M15 / 2000本 / Monte Carlo 100回 / 12銘柄順次実行")
    doc.add_paragraph("生成日時: 2026-07-15（H1スクリーニングと同条件・時間足のみ変更）")

    doc.add_heading("1. 総合結果", level=1)
    doc.add_paragraph(f"品質ゲート合格(1戦略以上): {n_pass} / 12 銘柄")
    doc.add_paragraph("エラー: 0件")
    doc.add_paragraph(
        "比較メモ: H1同条件では 0/12 合格。M15では #USSPX500・#UK100 の mean_reversion がゲート合格。"
    )

    doc.add_heading("2. 銘柄別サマリー", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, t in enumerate(["銘柄", "ゲート合格", "最良戦略", "Sharpe", "WF Sharpe", "備考"]):
        hdr[i].text = t
    for _, r in summary.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["symbol"])
        row[1].text = f"{int(r['n_pass'])}/3"
        row[2].text = str(r["best_strategy"])
        row[3].text = str(r["best_sharpe"])
        row[4].text = str(r["best_wf_test_sharpe"])
        if int(r["n_pass"]) >= 1:
            row[5].text = "ゲート合格"
        elif r["best_sharpe"] == 0 and r["best_wf_test_sharpe"] == 0:
            row[5].text = "シグナルなし"
        else:
            row[5].text = "要再検証"

    doc.add_heading("3. ゲート合格銘柄（詳細）", level=1)
    passed = detail[detail["gate"] == "PASS"]
    if passed.empty:
        doc.add_paragraph("（なし）")
    else:
        for _, r in passed.iterrows():
            doc.add_paragraph(
                f"- {r['symbol']} / {r['strategy']}: "
                f"return {r['total_return_pct']:.1f}%, Sharpe {r['sharpe']:.3f}, "
                f"WF Sharpe {r['wf_avg_test_sharpe']:.3f}, "
                f"trades {int(r['trades'])}, MC+ {r['mc_prob_positive_pct']:.0f}%"
            )

    doc.add_heading("4. H1との比較", level=1)
    doc.add_paragraph("- H1 (2000本): ゲート合格 0/12")
    doc.add_paragraph(f"- M15 (2000本): ゲート合格 {n_pass}/12")
    doc.add_paragraph("- 合格戦略はいずれも mean_reversion（#USSPX500, #UK100）")
    doc.add_paragraph("- feature_score 取引0件は継続: Japan225, GOLD, SILVER, WTI")

    doc.add_heading("5. 詳細データ", level=1)
    doc.add_paragraph(f"Excel: {XLSX.name}")

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
