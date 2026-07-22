from __future__ import annotations

import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

import pandas as pd
from docx import Document

from scripts.screening_params import BEST_MR_RSI_BY_SYMBOL, rsi_band_label

BASE = PROJECT_ROOT / "data/reports"

BEST_RSI = {symbol: rsi_band_label(symbol) for symbol in BEST_MR_RSI_BY_SYMBOL}

RSI_FILES = {
    "30/70": BASE / "screening_M30_2000bars_nogate_20260715_133134.xlsx",
    "25/75": BASE / "screening_M30_2000bars_nogate_rsi2575_20260715_201256.xlsx",
    "35/65": BASE / "screening_M30_2000bars_nogate_rsi3565_20260715_204330.xlsx",
}

NEW = BASE / "screening_M30_2000bars_nogate_rsibest_bb2080_20260715_213606.xlsx"
DOCX = NEW.with_name(NEW.stem + "_report.docx")


def mr_rows(path: Path) -> pd.DataFrame:
    d = pd.read_excel(path, sheet_name="StrategyDetail")
    return d[d["strategy"] == "mean_reversion"].copy()


def main() -> None:
    new = mr_rows(NEW)
    baselines: list[dict] = []
    for symbol, band in BEST_RSI.items():
        src = mr_rows(RSI_FILES[band])
        row = src[src["symbol"] == symbol].iloc[0]
        baselines.append(
            {
                "symbol": symbol,
                "rsi": band,
                "ret_bb15": float(row["total_return_pct"]),
                "sharpe_bb15": float(row["sharpe"]),
                "mdd_bb15": float(row["max_drawdown_pct"]),
                "trades_bb15": int(row["trades"]),
                "mc_bb15": float(row["mc_prob_positive_pct"]),
            }
        )
    base_df = pd.DataFrame(baselines)
    cmp = base_df.merge(
        new[
            [
                "symbol",
                "total_return_pct",
                "sharpe",
                "max_drawdown_pct",
                "trades",
                "mc_prob_positive_pct",
            ]
        ].rename(
            columns={
                "total_return_pct": "ret_bb20",
                "sharpe": "sharpe_bb20",
                "max_drawdown_pct": "mdd_bb20",
                "trades": "trades_bb20",
                "mc_prob_positive_pct": "mc_bb20",
            }
        ),
        on="symbol",
    )
    cmp["ret_delta"] = cmp["ret_bb20"] - cmp["ret_bb15"]
    cmp["sharpe_delta"] = cmp["sharpe_bb20"] - cmp["sharpe_bb15"]
    cmp = cmp.sort_values("ret_delta", ascending=False)

    print("=== Best RSI fixed | BB 0.15/0.85 -> 0.20/0.80 ===")
    print(
        cmp[
            [
                "symbol",
                "rsi",
                "ret_bb15",
                "ret_bb20",
                "ret_delta",
                "sharpe_bb15",
                "sharpe_bb20",
                "mdd_bb15",
                "mdd_bb20",
                "trades_bb15",
                "trades_bb20",
            ]
        ].to_string(index=False)
    )

    doc = Document()
    doc.add_heading("M30 BB entry 0.20/0.80 全銘柄テスト", 0)
    doc.add_paragraph("条件: M30 / 2000本 / MC100 / ゲートOFF")
    doc.add_paragraph("RSI: 銘柄別最良値に固定")
    doc.add_paragraph("変更: bb_entry 0.15/0.85 → 0.20/0.80")

    doc.add_heading("1. mean_reversion 比較", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for i, t in enumerate(["銘柄", "RSI", "BB0.15 return", "BB0.20 return", "Δ", "判定"]):
        table.rows[0].cells[i].text = t
    for _, r in cmp.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["symbol"])
        row[1].text = str(r["rsi"])
        row[2].text = f"{r['ret_bb15']:.1f}%"
        row[3].text = f"{r['ret_bb20']:.1f}%"
        row[4].text = f"{r['ret_delta']:+.1f}%"
        if r["ret_delta"] > 0.3:
            row[5].text = "改善"
        elif r["ret_delta"] < -0.3:
            row[5].text = "悪化"
        else:
            row[5].text = "ほぼ同等"

    improved = int((cmp["ret_delta"] > 0.3).sum())
    worse = int((cmp["ret_delta"] < -0.3).sum())
    same = len(cmp) - improved - worse
    doc.add_heading("2. 所見", level=1)
    doc.add_paragraph(f"改善: {improved} / 悪化: {worse} / ほぼ同等: {same}")
    doc.add_paragraph(
        "BB 0.20/0.80 は銘柄により効果が分かれる。"
        "改善銘柄のみ採用し、それ以外は 0.15/0.85 維持が妥当。"
    )
    doc.add_paragraph(f"Excel: {NEW.name}")
    doc.save(DOCX)
    print("Word:", DOCX)


if __name__ == "__main__":
    main()
