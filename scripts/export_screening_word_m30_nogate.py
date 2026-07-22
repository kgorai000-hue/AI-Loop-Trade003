from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document

XLSX = Path(
    r"C:\Users\kgora\Documents\MT5_loop\data\reports"
    r"\screening_M30_2000bars_nogate_20260715_133134.xlsx"
)
DOCX = XLSX.with_suffix(".docx")


def main() -> None:
    summary = pd.read_excel(XLSX, sheet_name="Summary")
    detail = pd.read_excel(XLSX, sheet_name="StrategyDetail")

    print("=== Summary ===")
    print(summary.to_string(index=False))
    print()

    traded = detail[detail["trades"] > 0].copy()
    traded = traded.sort_values(["total_return_pct", "sharpe"], ascending=[False, False])
    print("=== Top by total_return (trades>0) ===")
    cols = [
        "symbol",
        "strategy",
        "total_return_pct",
        "sharpe",
        "wf_avg_test_sharpe",
        "max_drawdown_pct",
        "trades",
        "mc_prob_positive_pct",
    ]
    print(traded[cols].head(12).to_string(index=False))

    doc = Document()
    doc.add_heading("MT5_loop 全銘柄スクリーニング（M30・品質ゲートOFF）", 0)
    doc.add_paragraph("条件: M30 / 2000本 / Monte Carlo 100回 / 品質ゲート無効 / 12銘柄順次")
    doc.add_paragraph("生成日時: 2026-07-15")

    doc.add_heading("1. 総合結果", level=1)
    doc.add_paragraph("品質ゲート: すべて無効（gate列は N/A）")
    doc.add_paragraph("エラー: 0件")
    doc.add_paragraph("評価方針: Sharpe / total return / WF などの実測値で相対比較")

    doc.add_heading("2. 銘柄別サマリー（最良Sharpe戦略）", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, t in enumerate(["銘柄", "最良戦略", "Sharpe", "WF Sharpe", "備考"]):
        table.rows[0].cells[i].text = t
    for _, r in summary.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["symbol"])
        row[1].text = str(r["best_strategy"])
        row[2].text = str(r["best_sharpe"])
        row[3].text = str(r["best_wf_test_sharpe"])
        row[4].text = "シグナルなし" if r["best_sharpe"] == 0 else "相対比較のみ"

    doc.add_heading("3. リターン上位（取引あり）", level=1)
    for _, r in traded.head(10).iterrows():
        line = (
            f"- {r['symbol']} / {r['strategy']}: "
            f"return {r['total_return_pct']:.1f}%, Sharpe {r['sharpe']:.3f}, "
            f"MDD {r['max_drawdown_pct']:.1f}%, WF {r['wf_avg_test_sharpe']:.3f}, "
            f"MC+ {r['mc_prob_positive_pct']:.0f}%"
        )
        doc.add_paragraph(line)

    doc.add_heading("4. 備考", level=1)
    doc.add_paragraph(
        "本結果はゲート無効のスクリーニング。"
        "最終採用にはゲートON・長期データでの再検証が必要。"
    )
    doc.add_paragraph(f"Excel: {XLSX.name}")

    doc.save(DOCX)
    print("Word:", DOCX)


if __name__ == "__main__":
    main()
