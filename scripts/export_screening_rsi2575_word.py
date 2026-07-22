from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]
NEW = PROJECT_ROOT / "data/reports/screening_M30_2000bars_nogate_rsi2575_20260715_201256.xlsx"
OLD = PROJECT_ROOT / "data/reports/screening_M30_2000bars_nogate_20260715_133134.xlsx"
DOCX = NEW.with_name(NEW.stem + "_report.docx")


def main() -> None:
    new = pd.read_excel(NEW, sheet_name="StrategyDetail")
    old = pd.read_excel(OLD, sheet_name="StrategyDetail")

    mr_new = new[new["strategy"] == "mean_reversion"].copy()
    mr_old = old[old["strategy"] == "mean_reversion"].copy()

    cmp = mr_old.merge(mr_new, on="symbol", suffixes=("_30_70", "_25_75"))
    cmp["return_delta"] = cmp["total_return_pct_25_75"] - cmp["total_return_pct_30_70"]
    cmp["sharpe_delta"] = cmp["sharpe_25_75"] - cmp["sharpe_30_70"]
    cmp = cmp.sort_values("return_delta", ascending=False)

    doc = Document()
    doc.add_heading("M30 mean_reversion RSI 25/75 全銘柄テスト", 0)
    doc.add_paragraph("条件: M30 / 2000本 / MC100 / ゲートOFF")
    doc.add_paragraph("変更: rsi_oversold=25, rsi_overbought=75（baseline 30/70）")

    doc.add_heading("1. 改善した銘柄（return ↑）", level=1)
    improved = cmp[cmp["return_delta"] > 0]
    for _, r in improved.iterrows():
        doc.add_paragraph(
            f"- {r['symbol']}: {r['total_return_pct_30_70']:.1f}% → "
            f"{r['total_return_pct_25_75']:.1f}% (Δ{r['return_delta']:+.1f}%), "
            f"Sharpe {r['sharpe_30_70']:.3f} → {r['sharpe_25_75']:.3f}"
        )

    doc.add_heading("2. 悪化した銘柄（return ↓）", level=1)
    worse = cmp[cmp["return_delta"] < 0].sort_values("return_delta")
    for _, r in worse.iterrows():
        doc.add_paragraph(
            f"- {r['symbol']}: {r['total_return_pct_30_70']:.1f}% → "
            f"{r['total_return_pct_25_75']:.1f}% (Δ{r['return_delta']:+.1f}%)"
        )

    doc.add_heading("3. 所見", level=1)
    doc.add_paragraph(
        "RSI 25/75 は米国指数（#US30, #USNDAQ100）でリターン・Sharpe が改善。"
        "#Japan225・WTI では大幅悪化。銘柄ごとに最適RSIが異なる。"
    )
    doc.add_paragraph(f"詳細Excel: {NEW.name}")
    doc.add_paragraph(f"比較baseline: {OLD.name}")

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
