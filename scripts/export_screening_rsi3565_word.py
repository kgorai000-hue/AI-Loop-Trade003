from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]
BASE = PROJECT_ROOT / "data/reports"
P_30 = BASE / "screening_M30_2000bars_nogate_20260715_133134.xlsx"
P_25 = BASE / "screening_M30_2000bars_nogate_rsi2575_20260715_201256.xlsx"
P_35 = BASE / "screening_M30_2000bars_nogate_rsi3565_20260715_204330.xlsx"
DOCX = P_35.with_name(P_35.stem + "_report.docx")


def mr_frame(path: Path, suffix: str) -> pd.DataFrame:
    d = pd.read_excel(path, sheet_name="StrategyDetail")
    m = d[d["strategy"] == "mean_reversion"][
        [
            "symbol",
            "total_return_pct",
            "sharpe",
            "wf_avg_test_sharpe",
            "max_drawdown_pct",
            "trades",
            "mc_prob_positive_pct",
        ]
    ].copy()
    return m.rename(columns={c: f"{c}_{suffix}" for c in m.columns if c != "symbol"})


def pick_best(row: pd.Series) -> str:
    candidates = [
        ("30/70", float(row["total_return_pct_30_70"])),
        ("25/75", float(row["total_return_pct_25_75"])),
        ("35/65", float(row["total_return_pct_35_65"])),
    ]
    return max(candidates, key=lambda x: x[1])[0]


def main() -> None:
    cmp = mr_frame(P_30, "30_70").merge(mr_frame(P_25, "25_75"), on="symbol")
    cmp = cmp.merge(mr_frame(P_35, "35_65"), on="symbol")
    cmp["best"] = cmp.apply(pick_best, axis=1)
    cmp = cmp.sort_values("total_return_pct_35_65", ascending=False)

    print("=== mean_reversion 3-way RSI comparison (return %) ===")
    print(
        cmp[
            [
                "symbol",
                "total_return_pct_30_70",
                "total_return_pct_25_75",
                "total_return_pct_35_65",
                "best",
            ]
        ].to_string(index=False)
    )
    print()
    print("=== best RSI count ===")
    print(cmp["best"].value_counts().to_string())

    doc = Document()
    doc.add_heading("M30 mean_reversion RSI 35/65 全銘柄テスト", 0)
    doc.add_paragraph("条件: M30 / 2000本 / MC100 / ゲートOFF")
    doc.add_paragraph("変更: rsi_oversold=35, rsi_overbought=65")

    doc.add_heading("1. 3方式 return 比較", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, t in enumerate(["銘柄", "30/70", "25/75", "35/65", "最良"]):
        table.rows[0].cells[i].text = t
    for _, r in cmp.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["symbol"])
        row[1].text = f"{r['total_return_pct_30_70']:.1f}%"
        row[2].text = f"{r['total_return_pct_25_75']:.1f}%"
        row[3].text = f"{r['total_return_pct_35_65']:.1f}%"
        row[4].text = str(r["best"])

    doc.add_heading("2. 所見", level=1)
    doc.add_paragraph(
        "銘柄ごとに最適RSIが異なる。"
        "全銘柄共通値は採用せず、銘柄別に最良バンドを選ぶのが妥当。"
    )
    for band, count in cmp["best"].value_counts().items():
        doc.add_paragraph(f"- 最良 {band}: {int(count)} 銘柄")
    doc.add_paragraph(f"Excel: {P_35.name}")

    doc.save(DOCX)
    print("Word:", DOCX)


if __name__ == "__main__":
    main()
