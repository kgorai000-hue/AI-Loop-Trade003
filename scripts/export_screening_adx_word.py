from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document

XLSX = Path(
    r"C:\Users\kgora\Documents\MT5_loop\data\reports"
    r"\screening_M30_2000bars_nogate_adxgrid_20260715_222307.xlsx"
)
DOCX = XLSX.with_name(XLSX.stem + "_report.docx")


def main() -> None:
    best = pd.read_excel(XLSX, sheet_name="BestPerSymbol")
    doc = Document()
    doc.add_heading("M30 ADX trend_threshold グリッド（18-30 step2）", 0)
    doc.add_paragraph("条件: M30 / 2000本 / MC100 / ゲートOFF / trend_followingのみ")

    doc.add_heading("1. 銘柄別最良ADX（return最大）", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["銘柄", "ADX", "return", "Sharpe", "MC+"]):
        table.rows[0].cells[i].text = h
    for _, r in best.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["symbol"])
        row[1].text = str(int(r["adx_trend_threshold"]))
        row[2].text = f"{r['total_return_pct']:.1f}%"
        row[3].text = f"{r['sharpe']:.3f}"
        row[4].text = f"{r['mc_prob_positive_pct']:.0f}%"

    doc.add_heading("2. 所見", level=1)
    doc.add_paragraph(
        "プラスリターンの最良ADXは #US30(26), SILVER(18), #Germany40(18), #Japan225(30) のみ。"
    )
    doc.add_paragraph("baseline 25 はグリッドに無いため、近傍は 24/26。")
    doc.add_paragraph(f"Excel: {XLSX.name}")
    doc.save(DOCX)
    print(DOCX)
    print(
        best[
            [
                "symbol",
                "adx_trend_threshold",
                "total_return_pct",
                "sharpe",
                "mc_prob_positive_pct",
            ]
        ].to_string(index=False)
    )


if __name__ == "__main__":
    main()
